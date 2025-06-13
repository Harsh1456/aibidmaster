from flask import Flask, request, jsonify, render_template, session, redirect, url_for, make_response
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
from flask_cors import CORS
from dotenv import load_dotenv
import math
import base64
import os
import re
import csv
import json
import pdfkit
from dateutil.parser import parse as parse_date
import PyPDF2
import openai
from openai import OpenAI
from docx import Document
from io import BytesIO, StringIO
from weasyprint import HTML
from weasyprint.text.fonts import FontConfiguration

load_dotenv()
app = Flask(__name__)
CORS(app)
app.secret_key = os.environ.get('SECRET_KEY')
client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))

app.config['ENV'] = os.environ.get('FLASK_ENV', 'production')

app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

ADMIN_EMAIL = os.environ.get('ADMIN_EMAIL', 'admin@aibidmaster.com')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'Admin@1235')

# Industry-standard constants
ASPHALT_DENSITY = 145  # lbs per cubic foot
ASPHALT_THICKNESS = 0.25  # feet (3 inches)
CONCRETE_DENSITY = 150  # lbs per cubic foot
CONCRETE_THICKNESS = 0.33  # feet (4 inches)
LABOR_RATE = 52.50  # dollars per hour
MATERIAL_MARKUP = 1.25  # 25% markup
EQUIPMENT_RATE_MULTIPLIER = 1.15
PROFIT_MARGIN = 0.18  # 18%
OVERHEAD_RATE = 0.12  # 12%
ASPHALT_COST_PER_TON = 135  # Washington DC market rate
CONCRETE_COST_PER_YD = 165  # Washington DC market rate

# Material unit costs (Washington DC market rates, 2024–2025)
MATERIAL_UNIT_COSTS = {
    'asphalt': 135,                # $/ton
    'concrete': 165,               # $/cubic yard
    'aggregate base': 38,          # $/ton
    'recycled asphalt': 110,       # $/ton
    'bituminous surface': 140,     # $/ton
    'subbase': 30,                 # $/ton
    'geotextile': 1.25,            # $/sq yd
    'emulsion': 3.5,               # $/gallon
    'sealcoat': 0.45,              # $/sq ft
    'thermoplastic striping': 2.5, # $/linear ft
    'curb': 38,                    # $/linear ft (concrete)
    'sidewalk': 12,                # $/sq ft (concrete)
    'pavers': 18,                  # $/sq ft
    'rebar': 0.65,                 # $/lb
    'drainage pipe': 38,           # $/linear ft
    'stormwater structure': 2500,  # $/each
    # Add more as needed
}

# Project Model 
class Project(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255), nullable=False)
    type = db.Column(db.String(100), nullable=False)
    location = db.Column(db.String(255), nullable=False)
    submitted = db.Column(db.Date, nullable=False)
    status = db.Column(db.Enum('pending', 'accepted', 'rejected'), default='pending')
    cost = db.Column(db.String(50), nullable=False)
    completion_date = db.Column(db.Date)
    land_mile = db.Column(db.Float)
    width = db.Column(db.Float)
    area = db.Column(db.Float, nullable=False)
    material = db.Column(db.String(50), nullable=False)
    tonnage = db.Column(db.Float)
    scope = db.Column(db.Text, nullable=False)
    requirements = db.Column(db.Text)
    estimated_cost = db.Column(db.String(50))
    profit_margin = db.Column(db.String(20))
    success_probability = db.Column(db.String(20))
    asphalt_tons = db.Column(db.Float)
    concrete_yds = db.Column(db.Float)
    rebar_lbs = db.Column(db.Float)
    aggregate_tons = db.Column(db.Float)
    management_hours = db.Column(db.Integer)
    prep_hours = db.Column(db.Integer)
    paving_hours = db.Column(db.Integer)
    finishing_hours = db.Column(db.Integer)

# Create tables before first request
# @app.before_first_request
# def create_tables():
#     db.create_all()

# Add root route to serve index.html
@app.route('/')
def index():
    return render_template('index.html')


# Admin routes
@app.route('/admin', methods=['GET'])
def admin_login_page():
    return render_template('admin_login.html')


@app.route('/admin/login', methods=['POST'])
def admin_login():
    data = request.form
    email = data.get('email')
    password = data.get('password')
    
    if email == ADMIN_EMAIL and password == ADMIN_PASSWORD:
        session['admin_logged_in'] = True
        return jsonify({'success': True})
    
    return jsonify({'success': False, 'message': 'Invalid credentials, Please Try Again'})


@app.route('/admin/dashboard', methods=['GET'])
def admin_dashboard():
    if not session.get('admin_logged_in'):
        return redirect(url_for('admin_login_page'))
    
    status = request.args.get('status', 'pending')
    query = Project.query
    
    if status != 'all':
        query = query.filter_by(status=status)
    
    projects = query.all()
    return render_template('admin_dashboard.html', 
                           status=status,
                           projects=projects)


@app.route('/admin/project/<int:project_id>', methods=['GET'])
def admin_project_detail(project_id):
    if not session.get('admin_logged_in'):
        return redirect(url_for('admin_login_page'))
    
    project = next((p for p in projects if p['id'] == project_id), None)
    if not project:
        return "Project not found", 404
    
    return render_template('admin_dashboard.html', project=project)


@app.route('/admin/logout', methods=['GET'])
def admin_logout():
    session.pop('admin_logged_in', None)
    return redirect(url_for('admin_login_page'))


@app.route('/api/admin/projects', methods=['GET'])
def get_projects():
    if not session.get('admin_logged_in'):
        return jsonify({'error': 'Unauthorized'}), 401
    
    status = request.args.get('status', 'pending')
    
    if status == 'all':
        projects = Project.query.all()
    else:
        projects = Project.query.filter_by(status=status).all()
    
    return jsonify([{
        'id': p.id,
        'name': p.name,
        'type': p.type,
        'location': p.location,
        'submitted': p.submitted.strftime('%Y-%m-%d'),
        'status': p.status,
        'cost': p.cost,
        # Include other fields as needed
    } for p in projects])


# Project Accept
@app.route('/api/admin/projects/<int:project_id>/accept', methods=['POST'])
def accept_project(project_id):
    project = Project.query.get(project_id)
    if project:
        project.status = 'accepted'
        db.session.commit()
        return jsonify({'message': 'Project accepted'})
    return jsonify({'error': 'Project not found'}), 404


# Project Reject
@app.route('/api/admin/projects/<int:project_id>/reject', methods=['POST'])
def reject_project(project_id):
    project = Project.query.get(project_id)
    if project:
        project.status = 'rejected'
        db.session.commit()
        return jsonify({'message': 'Project rejected'})
    return jsonify({'error': 'Project not found'}), 404


# Project Delete
@app.route('/api/admin/projects/<int:project_id>', methods=['DELETE'])
def delete_project(project_id):
    project = Project.query.get(project_id)
    if project:
        db.session.delete(project)
        db.session.commit()
        return jsonify({'message': 'Project deleted'})
    return jsonify({'error': 'Project not found'}), 404


@app.route('/api/admin/projects/<int:project_id>', methods=['GET'])
def get_project(project_id):
    project = Project.query.get(project_id)
    if project:
        return jsonify({
            'id': project.id,
            'name': project.name,
            'type': project.type,
            'location': project.location,
            'submitted': project.submitted.strftime('%Y-%m-%d'),
            'status': project.status,
            'cost': project.cost,
            'details': {
                'completionDate': project.completion_date.strftime('%Y-%m-%d') if project.completion_date else None,
                'landMile': project.land_mile,
                'width': project.width,
                'area': project.area,
                'material': project.material,
                'tonnage': project.tonnage,
                'scope': project.scope,
                'requirements': project.requirements,
                'estimatedCost': project.estimated_cost,
                'profitMargin': project.profit_margin,
                'successProbability': project.success_probability,
                'asphalt': project.asphalt_tons,
                'concrete': project.concrete_yds,
                'rebar': project.rebar_lbs,
                'aggregate': project.aggregate_tons,
                'managementHours': project.management_hours,
                'prepHours': project.prep_hours,
                'pavingHours': project.paving_hours,
                'finishingHours': project.finishing_hours
            }
        })
    return jsonify({'error': 'Project not found'}), 404


# PDF Processing
def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file):
    doc = Document(BytesIO(file.read()))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text += cell_text + "\n"
    if not text.strip():
        raise ValueError("No text found in DOCX.")
    return text

def extract_rfp_data(text):
    """Improved extraction with better patterns and fallbacks"""
    data = {}
    
    # Normalize text for easier matching
    normalized_text = re.sub(r'\s+', ' ', text).lower()
    
    # Extraction patterns - ordered by priority
    patterns = [
        # Project name/title
        (r'(?:project\s*name|project\s*title|project):?\s*([^\n]+)', 'project_name'),
        (r'project\s*#?\s*[\w-]+\s*-\s*([^\n]+)', 'project_name'),
        
        # Location
        (r'(?:project\s*location|location):?\s*([^\n]+)', 'project_location'),
        (r'in\s*([^\n,]+)(?:\s*county|\s*city|\s*state)', 'project_location'),
        
        # Dates
        (r'(?:completion\s*date|target\s*date|work\s*must\s*be\s*completed\s*by):?\s*([a-z]+\s\d{1,2},\s\d{4}|\d{4}-\d{2}-\d{2})', 'completion_date'),
        (r'fully\s*completed\s*by\s*([a-z]+\s\d{1,2},\s\d{4}|\d{4}-\d{2}-\d{2})', 'completion_date'),
        
        # Duration
        (r'(?:duration|project\s*duration)\s*\(?\s*weeks?\s*\)?:\s*(\d+)', 'project_duration'),
        
        # Measurements
        (r'(\d+(?:\.\d+)?)\s*lane\s*[-]?\s*mi(?:le)?s?', 'land_mile'),
        (r'(\d+(?:\.\d+)?)\s*(?:ft|feet|foot)(?:\s*width)?', 'width'),
        (r'area\s*\(?\s*sq\s*ft\s*\)?:\s*([\d,]+(?:\.\d+)?)', 'project_area'),
        (r'square\s*footage:\s*([\d,]+(?:\.\d+)?)', 'project_area'),
        
        # Material
        (r'\b(asphalt|hma|wma|concrete|aggregate\s*base)\b', 'material_type'),
        (r'tonnage:\s*([\d,]+(?:\.\d+)?)', 'tonnage'),
        (r'estimated\s*quantity:\s*([\d,]+(?:\.\d+)?)\s*tons?', 'tonnage'),
    ]
    
    # Apply patterns
    for pattern, key in patterns:
        if key not in data:  # Only capture first match for each field
            match = re.search(pattern, normalized_text, re.IGNORECASE)
            if match:
                data[key] = match.group(1).strip()
                if key in ['land_mile', 'width', 'tonnage']:
                    # Clean numeric values
                    data[key] = data[key].replace(',', '')
    
    # Special handling for area calculation
    if 'project_area' not in data and 'land_mile' in data and 'width' in data:
        try:
            land_mile = float(data['land_mile'])
            width = float(data['width'])
            data['project_area'] = str(round(land_mile * 5280 * width))
        except (ValueError, TypeError):
            pass
    
    # Extract scope and requirements sections
    section_patterns = [
        ('project_scope', r'scope\s*of\s*work:?'),
        ('project_requirements', r'(?:special\s*conditions|notes|special\s*requirements):?')
    ]
    
    for key, pattern in section_patterns:
        if key not in data:
            match = re.search(pattern, normalized_text, re.IGNORECASE)
            if match:
                start_pos = match.end()
                # Find the end of the section (next heading or blank line)
                end_pos = len(normalized_text)
                for end_pattern in [r'\n\s*\n', r'\n[A-Z][A-Z\s]+:']:
                    end_match = re.search(end_pattern, normalized_text[start_pos:], re.IGNORECASE)
                    if end_match:
                        end_pos = min(end_pos, start_pos + end_match.start())
                
                section_text = text[start_pos:end_pos].strip()
                if section_text:
                    data[key] = section_text
    
    # Clean extracted values
    for key in data:
        if isinstance(data[key], str):
            # Remove common trailing punctuation
            data[key] = re.sub(r'^[:;,.]+|[:;,.]+$', '', data[key].strip())
            # Capitalize first letter for certain fields
            if key in ['project_name', 'project_location', 'material_type']:
                data[key] = data[key][0].upper() + data[key][1:]
    
    return data


def extract_fields_with_openai(text):
    prompt = """
You are an expert at extracting structured data from construction RFPs. 
Given the following RFP text, extract and map all relevant fields to this schema, even if the field names in the RFP are different or in a different format. 
If a field is missing, use null or an empty string. 
If you find synonymous fields (e.g., "Job Title" for "project_name", "Place" for "project_location"), map them accordingly.

Respond with a JSON object with these keys:
- project_name
- project_type
- project_location
- completion_date
- project_duration
- land_mile
- width
- project_area
- material_type
- tonnage
- project_scope
- project_requirements

Text:
\"\"\"%s\"\"\"
""" % text[:3500]  # Send up to 3500 chars for context

    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        content = response.choices[0].message.content.strip()
        # Try to extract JSON from code block if present
        json_match = re.search(r'```json\n(.*?)\n```', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)
        data = json.loads(content)
        return data
    except Exception as e:
        app.logger.error(f"OpenAI GPT extraction failed: {str(e)}\nRaw output: {content if 'content' in locals() else ''}")
        return {}

@app.route('/upload_rfp', methods=['POST'])
def upload_rfp():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    try:
        filename = file.filename.lower()
        file_data = file.read()
        file_stream = BytesIO(file_data)
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_stream)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file_stream)
        else:
            return jsonify({'error': 'Unsupported file type'}), 400

        # Always use OpenAI extraction for robustness
        extracted_data = extract_fields_with_openai(text)

        # Set defaults for any missing required fields
        if not extracted_data.get('project_name'):
            extracted_data['project_name'] = f"Project from {file.filename}"
        if not extracted_data.get('project_type'):
            extracted_data['project_type'] = 'road'
        if not extracted_data.get('project_location'):
            extracted_data['project_location'] = 'Unknown Location'
        if not extracted_data.get('project_scope'):
            extracted_data['project_scope'] = 'Scope not extracted'

        # Calculate area if land-mile/width provided but no area
        if not extracted_data.get('project_area') and extracted_data.get('land_mile') and extracted_data.get('width'):
            try:
                land_mile = float(extracted_data['land_mile'])
                width_ft = float(extracted_data['width'])
                if land_mile > 0 and width_ft > 0:
                    extracted_data['project_area'] = str(round(land_mile * 5280 * width_ft))
            except (ValueError, TypeError):
                pass

        # Validate we have area
        if not extracted_data.get('project_area'):
            return jsonify({
                'error': 'Could not determine project area. Please provide area or land-mile+width in the document.',
                'extracted_data': extracted_data  # For debugging
            }), 400

        # Prepare data for processing
        data = {
            'project_name': extracted_data.get('project_name'),
            'project_type': extracted_data.get('project_type', 'road'),
            'project_location': extracted_data.get('project_location'),
            'project_duration': extracted_data.get('project_duration', ''),
            'completion_date': extracted_data.get('completion_date', ''),
            'land_mile': extracted_data.get('land_mile', ''),
            'width': extracted_data.get('width', ''),
            'project_area': extracted_data.get('project_area'),
            'material_type': extracted_data.get('material_type', 'asphalt'),
            'tonnage': extracted_data.get('tonnage', ''),
            'project_scope': extracted_data.get('project_scope'),
            'project_requirements': extracted_data.get('project_requirements', '')
        }

        # Pass to calculation phase
        return process_estimate(data)

    except Exception as e:
        app.logger.error(f"RFP processing failed: {str(e)}")
        return jsonify({
            'error': 'RFP processing failed',
            'details': str(e)
        }), 500


def process_estimate(data):
    app.logger.info(f"Starting estimate processing with data: {data}")

    def safe_float(value, default=0.0):
        if not value or not str(value).strip():
            return default
        # Remove commas and non-numeric (except dot and minus) characters
        cleaned = re.sub(r'[^\d\.\-]', '', str(value))
        try:
            return float(cleaned)
        except ValueError:
            return default

    try:
        # Extract project details
        project_name = data.get('project_name', 'Unnamed Project')
        project_type = data.get('project_type', 'road')
        location = data.get('project_location', 'Unknown Location')
        scope = data.get('project_scope', '')
        if not scope or not scope.strip():
            scope = 'Scope not provided'
        project_requirements = data.get('project_requirements', '')
        if not project_requirements:
            project_requirements = ''
        material_type = data.get('material_type', 'asphalt')
        # Truncate material_type to 50 chars to fit DB column
        material = db.Column(db.String(150), nullable=False)
        tonnage = safe_float(data.get('tonnage'))
        
        land_mile = safe_float(data.get('land_mile'))
        width_ft = safe_float(data.get('width'))
        area_sqft = safe_float(data.get('project_area'))
        
        # Calculate area if land-mile/width provided
        if area_sqft <= 0:
            if land_mile > 0 and width_ft > 0:
                area_sqft = land_mile * 5280 * width_ft
            else:
                return jsonify({
                    'error': 'Valid area required: Provide either area or land-mile+width'
                }), 400
        
        # Validate we have valid area
        if area_sqft <= 0:
            app.logger.error(f"Invalid area calculation: land_mile={land_mile}, width={width_ft}, area_sqft={area_sqft}")
            return jsonify({
                'error': 'Valid area required: Provide either area or land-mile+width',
                'details': f"land_mile: {land_mile}, width: {width_ft}, calculated_area: {area_sqft}"
            }), 400


        # Handle completion date and duration
        completion_date_str = data.get('completion_date', '')
        duration_weeks = safe_float(data.get('project_duration', '0'))
        
        if completion_date_str:
            try:
                completion_date = datetime.strptime(completion_date_str, '%Y-%m-%d')
                # Calculate duration based on completion date
                today = datetime.now()
                duration_weeks = max((completion_date - today).days / 7, 1)
            except:
                completion_date = datetime.now() + timedelta(weeks=8)
                duration_weeks = 8
        else:
            if duration_weeks <= 0:
                duration_weeks = 8
            completion_date = datetime.now() + timedelta(weeks=duration_weeks)
        
        # Material calculations
        material_estimates = calculate_materials(
            area_sqft, 
            material_type, 
            tonnage
        )
        
        # Labor calculations
        labor_estimates = calculate_labor(area_sqft, duration_weeks, project_type)
        
        # Equipment calculations
        equipment_estimates = calculate_equipment(area_sqft, duration_weeks)
        
        # Financial calculations
        financial_summary = calculate_financials(
            material_estimates, 
            labor_estimates, 
            equipment_estimates,
            area_sqft,
            duration_weeks
        )
        
        # Project summary
        project_summary = {
            'project_name': project_name,
            'project_type': project_type.capitalize(),
            'location': location,
            'completion_date': completion_date.strftime('%Y-%m-%d'),
            'duration_weeks': duration_weeks,
            'area_sqft': round(area_sqft),
            'material_type': material_type.capitalize()
        }
        
        # Success probability based on project factors
        success_probability = calculate_success_probability(project_type, area_sqft, duration_weeks)
        
        # Save project to database
        new_project = Project(
            name=project_name,
            type=project_type.capitalize(),
            location=location,
            submitted=datetime.now().date(),
            status='pending',
            cost=f"${financial_summary['total_cost']}",
            completion_date=completion_date.date(),
            land_mile=land_mile,
            width=width_ft,
            area=area_sqft,
            material=material_type.capitalize(),
            tonnage=tonnage if tonnage > 0 else material_estimates.get('asphalt_tons', material_estimates.get('concrete_yds', 0)),
            scope=scope,
            requirements=project_requirements,
            estimated_cost=f"${financial_summary['total_cost']}",
            profit_margin=financial_summary['profit_margin'],
            success_probability=success_probability,
            asphalt_tons=material_estimates.get('asphalt_tons', 0),
            concrete_yds=material_estimates.get('concrete_yds', 0),
            rebar_lbs=material_estimates.get('rebar_lbs', 0),
            aggregate_tons=material_estimates.get('aggregate_tons', 0),
            management_hours=labor_estimates['management_hours'],
            prep_hours=labor_estimates['prep_hours'],
            paving_hours=labor_estimates['paving_hours'],
            finishing_hours=labor_estimates['finishing_hours']
        )
        
        try:
            db.session.add(new_project)
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Database error details: {str(e)}")
            # Print the exact SQL error if available
            import traceback
            app.logger.error(traceback.format_exc())
            return jsonify({
                'error': 'Database operation failed',
                'details': str(e)
            }), 500

        # Prepare response
        response = {
            'project_summary': project_summary,
            'material_estimates': material_estimates,
            'labor_estimates': labor_estimates,
            'equipment_estimates': equipment_estimates,
            'financial_summary': financial_summary,
            'success_probability': success_probability,
            'project_id': new_project.id  # Make sure this is included
        }
        
        return jsonify(response), 200
    
    except Exception as e:
        app.logger.error(f"Estimate calculation failed: {str(e)}")
        # RETURN 500 WITH ERROR DETAILS
        return jsonify({
            'error': 'Estimate calculation failed',
            'details': str(e)
        }), 500

# Existing estimate calculation endpoint
@app.route('/calculate_estimate', methods=['POST'])
def calculate_estimate():
    data = request.json
    return process_estimate(data)

def calculate_materials(area_sqft, material_type, tonnage):
    """Calculate materials for various pavement types using DC rates and specs."""

    # Constants (define these at the module level)
    ASPHALT_THICKNESS = 0.33  # default 4 inches
    ASPHALT_DENSITY = 145     # lbs per cubic foot
    CONCRETE_THICKNESS = 0.5  # default 6 inches

    material_type = material_type.lower().strip()
    results = {}

    if material_type in ['asphalt', 'bituminous surface', 'recycled asphalt']:
        if tonnage > 0:
            asphalt_tons = tonnage
        else:
            thickness = ASPHALT_THICKNESS
            if material_type == 'bituminous surface':
                thickness = 0.17  # 2 inches
            elif material_type == 'recycled asphalt':
                thickness = 0.25  # 3 inches
            volume_cf = area_sqft * thickness
            asphalt_tons = (volume_cf * ASPHALT_DENSITY) / 2000
        
        results['asphalt_tons'] = round(asphalt_tons, 1)
        results['aggregate_tons'] = round(asphalt_tons * 1.25, 1)  # DC-adjusted base material
        results['rebar_lbs'] = round(area_sqft * 0.6)
        results['emulsion_gal'] = round(area_sqft * 0.06, 1)
        results['sealcoat_sqft'] = round(area_sqft)
        results['thermoplastic_strip_ft'] = round(area_sqft / 10)

    elif material_type in ['concrete', 'sidewalk', 'pavers']:
        thickness = CONCRETE_THICKNESS
        if material_type == 'sidewalk':
            thickness = 0.33  # 4 inches
        elif material_type == 'pavers':
            thickness = 0.17  # 2 inches

        volume_cf = area_sqft * thickness
        concrete_yds = volume_cf / 27
        results['concrete_yds'] = round(concrete_yds, 1)
        results['rebar_lbs'] = round(area_sqft * 1.4)
        results['aggregate_tons'] = round(concrete_yds * 1.6 * 1.25, 1)
        results['formwork_sqft'] = round(area_sqft * 1.15)
        if material_type == 'pavers':
            results['pavers_sqft'] = round(area_sqft)

    elif material_type == 'aggregate base':
        thickness = 0.5  # 6 inches
        volume_cf = area_sqft * thickness
        aggregate_tons = (volume_cf * 110) / 2000
        results['aggregate_tons'] = round(aggregate_tons, 1)

    elif material_type == 'subbase':
        thickness = 0.67  # 8 inches
        volume_cf = area_sqft * thickness
        subbase_tons = (volume_cf * 110) / 2000
        results['subbase_tons'] = round(subbase_tons, 1)

    elif material_type == 'geotextile':
        results['geotextile_sqyd'] = round(area_sqft / 9, 1)

    elif material_type == 'sealcoat':
        results['sealcoat_sqft'] = round(area_sqft)

    elif material_type == 'thermoplastic striping':
        results['thermoplastic_strip_ft'] = round(area_sqft / 10)

    elif material_type == 'curb':
        results['curb_ft'] = round(area_sqft / 5)

    elif material_type == 'drainage pipe':
        results['drainage_pipe_ft'] = round(area_sqft / 100)

    elif material_type == 'stormwater structure':
        results['stormwater_structures'] = max(1, round(area_sqft / 20000))

    else:
        # Default fallback to asphalt
        return calculate_materials(area_sqft, 'asphalt', tonnage)

    return results


# Updated labor calculations
def calculate_labor(area_sqft, duration_weeks, project_type):
    # DC-specific labor rates and productivity
    if "road" in project_type.lower():
        base_hours = 60  # hours per 1000 sq ft (includes prep, paving, finishing)
    else:
        base_hours = 75  # hours per 1000 sq ft for non-road projects
    
    # Add 20% premium for DC urban projects
    total_hours = (area_sqft / 1000) * base_hours * 1.20
    
    # Distribute hours across phases
    management_hours = total_hours * 0.12
    prep_hours = total_hours * 0.30
    paving_hours = total_hours * 0.40
    finishing_hours = total_hours * 0.18
    
    return {
        'management_hours': round(management_hours),
        'prep_hours': round(prep_hours),
        'paving_hours': round(paving_hours),
        'finishing_hours': round(finishing_hours),
        'total_hours': round(total_hours)
    }

def calculate_equipment(area_sqft, duration_weeks):
    # Calculate equipment needs based on DC productivity standards
    pavers = max(1, math.ceil(area_sqft / 80000))  # 1 paver per 80,000 sq ft
    rollers = max(1, math.ceil(area_sqft / 40000))  # 1 roller per 40,000 sq ft
    excavators = 1 if area_sqft < 150000 else 2
    trucks = max(2, math.ceil(area_sqft / 30000))  # 1 truck per 30,000 sq ft
    
    # Updated DC equipment rental rates ($/week)
    paver_cost = pavers * 2200 * duration_weeks
    roller_cost = rollers * 950 * duration_weeks
    excavator_cost = excavators * 1800 * duration_weeks
    truck_cost = trucks * 850 * duration_weeks
    
    return {
        'pavers': pavers,
        'rollers': rollers,
        'excavators': excavators,
        'trucks': trucks,
        'paver_cost': round(paver_cost),
        'roller_cost': round(roller_cost),
        'excavator_cost': round(excavator_cost),
        'truck_cost': round(truck_cost)
    }

def calculate_financials(materials, labor, equipment, area_sqft, duration_weeks):
    """Calculate financial summary based on all estimates and material types"""
    material_costs = 0

    # Loop through all material keys and sum costs
    for key, qty in materials.items():
        if key.endswith('_tons'):
            base = 'asphalt' if 'asphalt' in key else 'aggregate base'
            unit_cost = MATERIAL_UNIT_COSTS.get(base, 100)
            material_costs += qty * unit_cost * MATERIAL_MARKUP
        elif key == 'concrete_yds':
            material_costs += qty * MATERIAL_UNIT_COSTS['concrete'] * MATERIAL_MARKUP
        elif key == 'rebar_lbs':
            material_costs += qty * MATERIAL_UNIT_COSTS['rebar'] / 100  # per 100 lbs
        elif key == 'emulsion_gal':
            material_costs += qty * MATERIAL_UNIT_COSTS['emulsion']
        elif key == 'sealcoat_sqft':
            material_costs += qty * MATERIAL_UNIT_COSTS['sealcoat']
        elif key == 'thermoplastic_strip_ft':
            material_costs += qty * MATERIAL_UNIT_COSTS['thermoplastic striping']
        elif key == 'curb_ft':
            material_costs += qty * MATERIAL_UNIT_COSTS['curb']
        elif key == 'sidewalk_sqft':
            material_costs += qty * MATERIAL_UNIT_COSTS['sidewalk']
        elif key == 'pavers_sqft':
            material_costs += qty * MATERIAL_UNIT_COSTS['pavers']
        elif key == 'geotextile_sqyd':
            material_costs += qty * MATERIAL_UNIT_COSTS['geotextile']
        elif key == 'drainage_pipe_ft':
            material_costs += qty * MATERIAL_UNIT_COSTS['drainage pipe']
        elif key == 'stormwater_structures':
            material_costs += qty * MATERIAL_UNIT_COSTS['stormwater structure']
        elif key == 'subbase_tons':
            material_costs += qty * MATERIAL_UNIT_COSTS['subbase']
        # Add more as needed

    # Labor costs
    labor_costs = labor['total_hours'] * LABOR_RATE
    
    # Equipment costs
    equipment_costs = (
        equipment['paver_cost'] + 
        equipment['roller_cost'] + 
        equipment['excavator_cost'] + 
        equipment['truck_cost']
    ) * EQUIPMENT_RATE_MULTIPLIER
    
    # Subtotal costs
    subtotal = material_costs + labor_costs + equipment_costs
    
    # Additional costs
    overhead = subtotal * OVERHEAD_RATE
    profit = subtotal * PROFIT_MARGIN
    
    # Total cost
    total_cost = subtotal + overhead + profit
    
    # Cost breakdown
    cost_breakdown = {
        'materials': round(material_costs),
        'labor': round(labor_costs),
        'equipment': round(equipment_costs),
        'overhead': round(overhead),
        'profit': round(profit)
    }
    
    # Cost per sq ft
    cost_per_sqft = total_cost / area_sqft if area_sqft > 0 else 0
    
    return {
        'total_cost': round(total_cost),
        'cost_per_sqft': round(cost_per_sqft, 2),
        'profit_margin': f"{PROFIT_MARGIN * 100}%",
        'cost_breakdown': cost_breakdown
    }

def calculate_success_probability(project_type, area_sqft, duration_weeks):
    """Calculate probability of bid success based on project factors"""
    base_prob = 70  # Base 70% probability
    
    # Adjust based on project type
    if project_type == 'road':
        base_prob += 5
    elif project_type == 'renovation':
        base_prob -= 3
    
    # Adjust based on project size
    if area_sqft > 100000:
        base_prob -= 10  # Large projects are more competitive
    elif area_sqft < 10000:
        base_prob += 5  # Small projects have less competition
    
    # Adjust based on duration
    if duration_weeks > 26:  # >6 months
        base_prob -= 8
    elif duration_weeks < 8:  # <2 months
        base_prob += 5
    
    # Ensure within bounds
    probability = max(50, min(90, base_prob))
    
    return f"{probability}%"


# Download Report  
@app.route('/download_report/<int:project_id>', methods=['GET'])
def download_report(project_id):
    project = Project.query.get(project_id)
    if not project:
        return jsonify({'error': 'Project not found'}), 404
    
    # Create PDF report
    pdf = generate_pdf_report(project)
    
    response = make_response(pdf)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=project_{project_id}_report.pdf'
    return response

@app.route('/download_report_csv/<int:project_id>', methods=['GET'])
def download_report_csv(project_id):
    project = Project.query.get(project_id)
    if not project:
        return jsonify({'error': 'Project not found'}), 404
    
    # Create CSV report
    si = StringIO()
    cw = csv.writer(si)
    
    # Write CSV headers
    cw.writerow(['Project Report', f'Project ID: {project_id}'])
    cw.writerow([])
    cw.writerow(['Field', 'Value'])
    cw.writerow(['Project Name', project.name])
    cw.writerow(['Project Type', project.type])
    cw.writerow(['Location', project.location])
    cw.writerow(['Submitted Date', project.submitted.strftime('%Y-%m-%d')])
    cw.writerow(['Status', project.status])
    cw.writerow(['Estimated Cost', project.cost])
    cw.writerow(['Completion Date', project.completion_date.strftime('%Y-%m-%d') if project.completion_date else ''])
    cw.writerow(['Area (sq ft)', project.area])
    cw.writerow(['Material', project.material])
    cw.writerow(['Asphalt (tons)', project.asphalt_tons])
    cw.writerow(['Concrete (yds)', project.concrete_yds])
    cw.writerow(['Rebar (lbs)', project.rebar_lbs])
    cw.writerow(['Aggregate (tons)', project.aggregate_tons])
    cw.writerow(['Management Hours', project.management_hours])
    cw.writerow(['Preparation Hours', project.prep_hours])
    cw.writerow(['Paving Hours', project.paving_hours])
    cw.writerow(['Finishing Hours', project.finishing_hours])
    cw.writerow(['Profit Margin', project.profit_margin])
    cw.writerow(['Success Probability', project.success_probability])
    cw.writerow(['Scope', project.scope])
    cw.writerow(['Requirements', project.requirements or ''])
    
    response = make_response(si.getvalue())
    response.headers['Content-Type'] = 'text/csv'
    response.headers['Content-Disposition'] = f'attachment; filename=project_{project_id}_report.csv'
    return response


def generate_pdf_report(project):
    # Get current date for report
    current_date = datetime.now().strftime('%B %d, %Y')
    
    # Read and encode logo (add your logo.png in static/images)
    logo_path = os.path.join(app.root_path, 'static', 'images', 'logo.png')
    logo_data = ""
    if os.path.exists(logo_path):
        with open(logo_path, "rb") as logo_file:
            logo_data = base64.b64encode(logo_file.read()).decode('utf-8')
    
    html_content = f"""
    <html>
    <head>
        <title>Project Report - {project.id}</title>
        <style>
            /* Professional styling */
            @page {{ size: A4; margin: 1.5cm; }}
            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; color: #333; line-height: 1.6; }}
            .header {{ border-bottom: 2px solid #3498db; padding-bottom: 15px; margin-bottom: 25px; }}
            h1 {{ color: #2c3e50; margin-bottom: 5px; }}
            h2 {{ color: #3498db; border-bottom: 1px solid #eee; padding-bottom: 8px; margin-top: 25px; }}
            .subtitle {{ color: #7f8c8d; font-size: 1.1rem; }}
            .project-info {{ background-color: #f8f9fa; border-radius: 8px; padding: 20px; margin-bottom: 30px; }}
            .grid-container {{ display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }}
            .section {{ margin-bottom: 25px; }}
            table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
            th {{ background-color: #3498db; color: white; text-align: left; padding: 12px; }}
            td {{ padding: 10px; border-bottom: 1px solid #eee; }}
            tr:nth-child(even) {{ background-color: #f9f9f9; }}
            .footer {{ margin-top: 40px; text-align: center; color: #7f8c8d; font-size: 0.9rem; }}
            .status-badge {{
                display: inline-block;
                padding: 3px 10px;
                border-radius: 12px;
                font-size: 0.85rem;
                font-weight: bold;
                margin-left: 10px;
            }}
            .status-pending {{ background-color: #f39c12; color: white; }}
            .status-accepted {{ background-color: #27ae60; color: white; }}
            .status-rejected {{ background-color: #e74c3c; color: white; }}
            .header-logo-container {{ 
                display: flex;
                justify-content: center; /* Horizontal centering */
                margin-bottom: 10px;
            }}
            
            .logo-img {{ 
                max-width: 100%; 
                height: auto; 
            }}
            .logo {{ font-weight: bold; font-size: 1.8rem; margin-bottom: 5px; }}
            .logo span {{ color: #f39c12; }}
        </style>
    </head>
    <body>
        <div class="header-logo-container">
            {"<img src='data:image/png;base64," + logo_data + "' class='logo-img'/>" if logo_data else ""}
        </div>

        <div class="header">
            <div class="logo">Bid<span>Master</span></div>
            <h1>Project Report: {project.name}</h1>
            <p class="subtitle">Generated on {current_date} | Project ID: {project.id}</p>
        </div>

        
        <div class="project-info">
            <div class="grid-container">
                <div>
                    <p><strong>Project Type:</strong> {project.type}</p>
                    <p><strong>Location:</strong> {project.location}</p>
                    <p><strong>Submitted:</strong> {project.submitted.strftime('%Y-%m-%d')}</p>
                    <p><strong>Status:</strong> 
                        <span class="status-badge status-{project.status}">{project.status.capitalize()}</span>
                    </p>
                </div>
                <div>
                    <p><strong>Estimated Cost:</strong> {project.cost}</p>
                    <p><strong>Completion Date:</strong> {project.completion_date.strftime('%Y-%m-%d') if project.completion_date else 'N/A'}</p>
                    <p><strong>Area:</strong> {project.area} sq ft</p>
                    <p><strong>Material:</strong> {project.material}</p>
                </div>
            </div>
        </div>
        
        <div class="section">
            <h2>Financial Summary</h2>
            <table>
                <tr>
                    <td>Profit Margin</td>
                    <td>{project.profit_margin}</td>
                    <td>Success Probability</td>
                    <td>{project.success_probability}</td>
                </tr>
            </table>
        </div>
        
        <div class="section">
            <h2>Resource Estimates</h2>
            <table>
                <tr>
                    <th>Material</th>
                    <th>Quantity</th>
                </tr>
                <tr>
                    <td>Asphalt</td>
                    <td>{project.asphalt_tons} tons</td>
                </tr>
                <tr>
                    <td>Concrete</td>
                    <td>{project.concrete_yds} cubic yards</td>
                </tr>
                <tr>
                    <td>Rebar</td>
                    <td>{project.rebar_lbs} lbs</td>
                </tr>
                <tr>
                    <td>Aggregate</td>
                    <td>{project.aggregate_tons} tons</td>
                </tr>
            </table>
        </div>
        
        <div class="section">
            <h2>Labor Estimates</h2>
            <table>
                <tr>
                    <th>Task</th>
                    <th>Hours</th>
                </tr>
                <tr>
                    <td>Management</td>
                    <td>{project.management_hours}</td>
                </tr>
                <tr>
                    <td>Preparation</td>
                    <td>{project.prep_hours}</td>
                </tr>
                <tr>
                    <td>Paving</td>
                    <td>{project.paving_hours}</td>
                </tr>
                <tr>
                    <td>Finishing</td>
                    <td>{project.finishing_hours}</td>
                </tr>
            </table>
        </div>
        
        <div class="section">
            <h2>Project Scope</h2>
            <p>{project.scope}</p>
        </div>
        
        <div class="section">
            <h2>Requirements</h2>
            <p>{project.requirements or 'No special requirements specified'}</p>
        </div>
        
        <div class="footer">
            <p>Generated by Paveiq BidMaster System</p>
            <p>&copy; {datetime.now().year} Paveiq. All rights reserved.</p>
        </div>
    </body>
    </html>
    """
    
    font_config = FontConfiguration()
    return HTML(string=html_content).write_pdf(font_config=font_config)


if __name__ == '__main__':
    # app.run(debug=True, port=5000)
    app.run(host='0.0.0.0', port=5000, debug=False)
